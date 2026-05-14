#!/usr/bin/env python3
"""Generate DOCX with the data-structure diagram and a plain-language description."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

t = doc.add_heading('Struktura danych aplikacji — widok użytkownika', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Medykamenty · plany produkcji, produkty, surowce, dostawcy')
r.italic = True
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(
    '/Users/wmankowski/medykamenty/diagrams/struktura_danych.jpg',
    width=Inches(6.5),
)

doc.add_page_break()

# ---------------- Section 1 ----------------
doc.add_heading('1. Katalog: dostawcy ↔ surowce / komponenty → produkt', level=1)

doc.add_paragraph(
    'Katalog to „magazyn wiedzy” aplikacji — opisuje, z czego składa się każdy '
    'produkt i od kogo można kupić poszczególne materiały. Cztery główne kartoteki:'
)

def bullet(label, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(label)
    r.bold = True
    p.add_run(' — ' + text)

bullet(
    'Dostawca',
    'firma, od której kupujesz. Trzyma e-mail, telefon, osobę kontaktową, warunki '
    'płatności oraz preferowany język korespondencji (PL/EN), w którym wysyłane będą '
    'zapytania ofertowe.',
)
bullet(
    'Surowiec',
    'składnik receptury (np. olej, ekstrakt). Ma jednostkę miary, MOQ, lead time, '
    'ostatnią cenę zakupu, listę dostawców (z których jeden jest preferowany) oraz '
    'flagę „dostarcza fabryka” — wtedy nie zamawiasz go sam.',
)
bullet(
    'Komponent opakowaniowy',
    'tuba, butelka, etykieta, kapsel, pompka, pipeta, kartonik, ulotka, karton '
    'zbiorczy, taśma, beczka, worek, konfekcja. Każdy z własnym MOQ, ceną i listą '
    'dostawców.',
)
bullet(
    'Produkt',
    'receptura: nazwa, SKU, pojemność (ml), gęstość (g/ml), MOQ w sztukach, koszt '
    'konfekcji, ewentualna masa na saszetki próbne. Receptura składa się z dwóch list: '
    'składniki (surowiec + procent masy) oraz opakowania (komponent + ilość sztuk '
    'na jeden produkt).',
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Relacje wiele-do-wielu: ')
r.bold = True
p.add_run(
    'ten sam surowiec może mieć kilku dostawców, a jeden dostawca dostarcza wiele '
    'surowców i komponentów. Aplikacja zna preferowanego dostawcę dla każdej '
    'pozycji i to do niego pójdą zapytania ofertowe.'
)

# ---------------- Section 2 ----------------
doc.add_heading('2. Przepływ produkcyjny: plan → magazyn → braki → e-maile', level=1)

doc.add_paragraph('Tu zaczyna się codzienna praca. Cztery powiązane dokumenty:')

bullet(
    'Plan produkcji',
    '„co chcemy wyprodukować”. Lista pozycji w dwóch postaciach: items (produkt '
    '+ liczba sztuk) oraz bulkMass (produkt + kilogramy luzem). Plan ma status '
    '(szkic / obliczony / zarchiwizowany), a po realizacji można wpisać, ile '
    'naprawdę udało się wyprodukować.',
)
bullet(
    'Stan magazynowy',
    'zaimportowany z Excela snapshot tego, co aktualnie jest w magazynie. Każdą '
    'pozycję ze stanu aplikacja stara się sama dopasować do surowca lub komponentu '
    'z katalogu (po nazwie i symbolu MP).',
)
bullet(
    'Raport braków',
    'obliczany dla wybranego planu: bierze receptury produktów z planu, mnoży przez '
    'ilości, odejmuje to, co jest na stanie, i pokazuje listę tego, czego brakuje. '
    'Pozycje są grupowane wg preferowanego dostawcy i uwzględniają MOQ przy '
    'sugerowanej ilości zamówienia.',
)
bullet(
    'Paczka e-maili RFQ',
    'wygenerowana z raportu braków: po jednym e-mailu na dostawcę, automatycznie w '
    'jego preferowanym języku, z listą pozycji do wyceny / zamówienia. Treść można '
    'opcjonalnie poprawić przez AI przed wysłaniem.',
)

# ---------------- Cheat sheet ----------------
doc.add_heading('3. Ściągawka relacji', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Co z czym'
hdr[1].text = 'Typ relacji'
hdr[2].text = 'Gdzie to widać w danych'
for c in hdr:
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True

rows = [
    ('Dostawca ↔ Surowiec', 'N : M (z preferowanym)', 'RawMaterial.supplierIds + preferredSupplierId'),
    ('Dostawca ↔ Komponent', 'N : M (z preferowanym)', 'PackagingComponent.supplierIds + preferredSupplierId'),
    ('Produkt → Surowce', '1 : N (z procentem)', 'Product.ingredients[]'),
    ('Produkt → Komponenty', '1 : N (z ilością/szt.)', 'Product.packaging[]'),
    ('Plan → Produkty', '1 : N (items + bulkMass)', 'ProductionPlan.items / bulkMass'),
    ('Plan → Raport braków', '1 : N (każde obliczenie)', 'ShortageReportEntry.planId'),
    ('Raport → Paczka e-maili', '1 : 1', 'EmailBatch.reportId'),
    ('Magazyn → Surowiec/Komponent', 'dopasowanie po nazwie/MP', 'StockRow.matchedRawMaterialId / matchedComponentId'),
]
for label, kind, where in rows:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = kind
    cells[2].text = where

out = '/Users/wmankowski/medykamenty/diagrams/struktura_danych.docx'
doc.save(out)
print('Saved:', out)
